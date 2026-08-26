import os
import io
import json
import re
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponse, FileResponse, Http404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_GET
from apps.accounts.decorators import role_required
from apps.tools.backup_utils import (
    get_db_statistics, create_database_backup, list_backups,
    restore_database_backup, delete_backup, get_backup_dir, get_db_path
)

from apps.academics.models import Classroom, AcademicYear
from apps.students.models import Student

import pypdf
from docx import Document
from openpyxl import Workbook


@login_required
def tools_hub(request):
    """
    Central Online Tools Hub displaying all categories and utilities.
    """
    total_classrooms = Classroom.objects.count()
    return render(request, 'tools/hub.html', {
        'page_title': 'មជ្ឈមណ្ឌលឧបករណ៍អនឡាញ (Online Tools Hub)',
        'total_classrooms': total_classrooms,
    })


@login_required
def pdf_merge_view(request):
    """
    PDF Merge & Organizer Tool: Merge multiple PDF documents, re-order, preview.
    """
    return render(request, 'tools/pdf_merge.html', {
        'page_title': 'បញ្ចូលឯកសារ PDF ចូលគ្នា (PDF Merge)',
    })


@login_required
def pdf_split_view(request):
    """
    PDF Split & Page Extractor: Extract specific pages or split PDF into pieces.
    """
    return render(request, 'tools/pdf_split.html', {
        'page_title': 'បំបែក & ទាញយកទំព័រ PDF (PDF Split & Extract)',
    })


@login_required
def pdf_to_word_excel_view(request):
    """
    PDF to Word (.docx) & PDF to Excel (.xlsx) / CSV Table Extractor.
    """
    return render(request, 'tools/pdf_to_word_excel.html', {
        'page_title': 'បំលែង PDF ទៅជា Word & Excel (PDF to Word & Excel Converter)',
    })


@login_required
def images_to_pdf_view(request):
    """
    Images to PDF Converter: Convert exam photos, homework, documents to single PDF.
    """
    return render(request, 'tools/images_to_pdf.html', {
        'page_title': 'បំលែងរូបភាពជាឯកសារ PDF (Images to PDF)',
    })


@login_required
def doc_scanner_view(request):
    """
    Smart Document & Paper Scanner (CamScanner-Style) with live camera capture & paper filters.
    """
    return render(request, 'tools/doc_scanner.html', {
        'page_title': 'ម៉ាស៊ីនស្កេនក្រដាស & ឯកសារ (Smart Document Scanner)',
    })


@login_required
def image_editor_view(request):
    """
    Studio Image Editor: Crop, Rotate, Filter, Draw, Annotate, Add School Watermark.
    """
    return render(request, 'tools/image_editor.html', {
        'page_title': 'កម្មវិធីកែសម្រួលរូបភាព (Studio Image Editor)',
    })


@login_required
def id_photo_maker_view(request):
    """
    Student & Teacher ID Photo Maker (4x6 & 3x4 cm) with background replacer & multi-photo sheet.
    """
    return render(request, 'tools/id_photo_maker.html', {
        'page_title': 'បង្កើតរូបថតកាតសិស្ស-គ្រូ 4x6 & 3x4 (ID Photo Maker)',
    })


@login_required
def image_compressor_view(request):
    """
    Batch Image Compressor & Format Converter (PNG, JPG, WEBP).
    """
    return render(request, 'tools/image_compressor.html', {
        'page_title': 'បង្រួម & បំលែងប្រភេទរូបភាព (Image Compressor & Converter)',
    })


@login_required
def qr_generator_view(request):
    """
    Advanced QR Code Generator for Links, WiFi, vCards, Telegram with school logo.
    """
    return render(request, 'tools/qr_generator.html', {
        'page_title': 'បង្កើត QR Code គ្រប់ប្រភេទ (Advanced QR Generator)',
    })


@login_required
def qr_scanner_view(request):
    """
    QR Code & Barcode Scanner via device camera or file upload.
    """
    return render(request, 'tools/qr_scanner.html', {
        'page_title': 'ស្កេន QR Code & Barcode (QR & Barcode Scanner)',
    })


@login_required
def khmer_number_converter_view(request):
    """
    Khmer Number to Words & Currency Spellout Converter.
    """
    return render(request, 'tools/khmer_number_converter.html', {
        'page_title': 'បំលែងលេខទៅជាអក្សរខ្មែរ (Khmer Number to Words)',
    })


@login_required
def text_analyzer_view(request):
    """
    Word Counter, Character Counter, Khmer Text Analyzer & Reading Time Estimator.
    """
    return render(request, 'tools/text_analyzer.html', {
        'page_title': 'រាប់ពាក្យ & វិភាគអត្ថបទ (Word Counter & Text Analyzer)',
    })


@login_required
def voice_typing_view(request):
    """
    Voice Typing & Speech-to-Text Dictation supporting Khmer and English.
    """
    return render(request, 'tools/voice_typing.html', {
        'page_title': 'វាយអត្ថបទតាមសំឡេង (Voice Typing & Speech to Text)',
    })


@login_required
def classroom_picker_view(request):
    """
    Interactive Classroom Lucky Draw Wheel, Random Name Picker, Team Splitter & Stopwatch.
    """
    active_year = AcademicYear.objects.filter(is_current=True).first()
    if active_year:
        classrooms = Classroom.objects.filter(academic_year=active_year).order_by('grade_level', 'name')
    else:
        classrooms = Classroom.objects.all().order_by('grade_level', 'name')
    return render(request, 'tools/classroom_picker.html', {
        'page_title': 'ចាប់ឆ្នោតសិស្ស & ចែកក្រុមរៀន (Classroom Lucky Draw & Team Builder)',
        'classrooms': classrooms,
        'active_year': active_year,
    })


@login_required
def calculator_converter_view(request):
    """
    Scientific Calculator & Universal Educational Unit Converter.
    """
    return render(request, 'tools/calculator_converter.html', {
        'page_title': 'ម៉ាស៊ីនគិតលេខ & បំលែងខ្នាត (Scientific Calculator & Unit Converter)',
    })


# --------------------------------------------------------------------------
# API Endpoints
# --------------------------------------------------------------------------

@login_required
@require_GET
def api_classroom_students(request, classroom_id):
    """
    API returning student names and information for classroom lucky draw & team builder.
    """
    try:
        classroom = get_object_or_404(Classroom, id=classroom_id)
        students = Student.objects.filter(
            classroom=classroom,
            status=Student.Status.ACTIVE
        ).order_by('khmer_name')

        data = [{
            'id': s.id,
            'student_id': s.student_id or f"S-{s.id}",
            'khmer_name': s.khmer_name,
            'latin_name': s.latin_name or '',
            'gender': s.get_gender_display(),
            'gender_code': s.gender,
        } for s in students]

        return JsonResponse({
            'success': True,
            'classroom': {
                'id': classroom.id,
                'name': classroom.name,
                'total_students': len(data),
            },
            'students': data,
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_POST
def api_pdf_merge(request):
    """
    Server-side fallback endpoint to merge uploaded PDF files.
    """
    try:
        files = request.FILES.getlist('pdf_files')
        if not files:
            return JsonResponse({'success': False, 'error': 'មិនមានឯកសារ PDF ត្រូវបានជ្រើសរើសឡើយ'}, status=400)

        writer = pypdf.PdfWriter()
        for f in files:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                writer.add_page(page)

        output_stream = io.BytesIO()
        writer.write(output_stream)
        output_stream.seek(0)

        response = HttpResponse(output_stream.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="Merged_SchoolSM_Document.pdf"'
        return response
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'បរាជ័យក្នុងការបញ្ចូល PDF: {str(e)}'}, status=500)


@login_required
def normalize_khmer_text(text):
    """
    Cleans up broken spacing between Khmer syllables and corrects misplaced pre-vowels.
    """
    if not text:
        return ""
    # Reorder pre-vowels (េ, ែ, ៃ) placed before consonants
    text = re.sub(r'([\u17C1\u17C2\u17C3])([\u1780-\u17B3])', r'\2\1', text)
    # Remove spaces between base consonants and coeng markers
    text = re.sub(r'\s+(\u17D2[\u1780-\u17B3])', r'\1', text)
    # Remove spaces between base consonants and dependent vowels
    text = re.sub(r'([\u1780-\u17B3])\s+([\u17B6-\u17C5])', r'\1\2', text)
    return text


@login_required
@require_POST
def api_pdf_to_docx(request):
    """
    Extract text and structured tables from uploaded PDF into Microsoft Word (.docx).
    Automatically detects table structures and builds real Word tables with borders and columns.
    """
    try:
        pdf_file = request.FILES.get('pdf_file')
        if not pdf_file:
            return JsonResponse({'success': False, 'error': 'សូមជ្រើសរើសឯកសារ PDF'}, status=400)

        reader = pypdf.PdfReader(pdf_file)
        doc = Document()

        # Set default font to Kantumruy Pro / Khmer OS Battambang
        for style in doc.styles:
            if hasattr(style, 'font'):
                style.font.name = 'Kantumruy Pro'

        original_name = pdf_file.name.rsplit('.', 1)[0] if '.' in pdf_file.name else 'Document'
        doc.add_heading(f'ឯកសារស្រង់ចេញ៖ {original_name}', level=1)

        total_pages = len(reader.pages)
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if not text:
                continue

            if total_pages > 1:
                doc.add_heading(f'--- ទំព័រទី {idx} ---', level=2)

            lines = [normalize_khmer_text(l.strip()) for l in text.split('\n') if l.strip()]
            
            # Group lines into paragraphs and table blocks
            current_table_rows = []

            for line in lines:
                # Detect if line is a table row (contains multiple spaces, tabs, or pipe symbols)
                parts = [p.strip() for p in re.split(r'\s{2,}|\t|\|', line) if p.strip()]

                if len(parts) >= 2:
                    current_table_rows.append(parts)
                else:
                    # If we had a table block accumulated, write it to docx as a real Table
                    if current_table_rows:
                        if len(current_table_rows) >= 2 or any(len(r) >= 3 for r in current_table_rows):
                            # Real Table detected
                            max_cols = max(len(r) for r in current_table_rows)
                            table = doc.add_table(rows=len(current_table_rows), cols=max_cols)
                            table.style = 'Table Grid'
                            for r_idx, row_data in enumerate(current_table_rows):
                                for c_idx, cell_value in enumerate(row_data):
                                    cell = table.cell(r_idx, c_idx)
                                    cell.text = cell_value
                        else:
                            for r in current_table_rows:
                                doc.add_paragraph("    ".join(r))
                        current_table_rows = []

                    # Add regular paragraph
                    doc.add_paragraph(line)

            # Flush any remaining table block
            if current_table_rows:
                if len(current_table_rows) >= 2 or any(len(r) >= 3 for r in current_table_rows):
                    max_cols = max(len(r) for r in current_table_rows)
                    table = doc.add_table(rows=len(current_table_rows), cols=max_cols)
                    table.style = 'Table Grid'
                    for r_idx, row_data in enumerate(current_table_rows):
                        for c_idx, cell_value in enumerate(row_data):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_value
                else:
                    for r in current_table_rows:
                        doc.add_paragraph("    ".join(r))

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        response = HttpResponse(
            output_stream.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{original_name}_converted.docx"'
        return response
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'បរាជ័យក្នុងការបំលែងជា Word: {str(e)}'}, status=500)


@login_required
@require_POST
def api_pdf_to_excel(request):
    """
    Extract text/tables from uploaded PDF into Microsoft Excel (.xlsx) with structured columns.
    """
    try:
        pdf_file = request.FILES.get('pdf_file')
        if not pdf_file:
            return JsonResponse({'success': False, 'error': 'សូមជ្រើសរើសឯកសារ PDF'}, status=400)

        reader = pypdf.PdfReader(pdf_file)
        wb = Workbook()
        ws = wb.active
        ws.title = "Extracted Data"

        row_num = 1
        ws.cell(row=row_num, column=1, value=f"ទិន្នន័យស្រង់ចេញពី PDF៖ {pdf_file.name}")
        row_num += 2

        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text:
                if len(reader.pages) > 1:
                    ws.cell(row=row_num, column=1, value=f"--- ទំព័រទី {idx} ---")
                    row_num += 1

                lines = text.split('\n')
                for line in lines:
                    line_str = line.strip()
                    if not line_str:
                        continue
                    parts = [p.strip() for p in re.split(r'\s{2,}|\t|\|', line_str) if p.strip()]
                    if parts:
                        for col_idx, part in enumerate(parts, start=1):
                            ws.cell(row=row_num, column=col_idx, value=part)
                        row_num += 1
                row_num += 1

        output_stream = io.BytesIO()
        wb.save(output_stream)
        output_stream.seek(0)

        original_name = pdf_file.name.rsplit('.', 1)[0] if '.' in pdf_file.name else 'Data'
        response = HttpResponse(
            output_stream.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{original_name}_extracted.xlsx"'
        return response
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'បរាជ័យក្នុងការបំលែងជា Excel: {str(e)}'}, status=500)


@login_required
@require_POST
def api_images_to_pdf(request):
    """
    Compile multiple uploaded images into a clean, high-quality PDF document.
    """
    try:
        images_files = request.FILES.getlist('images')
        if not images_files:
            return JsonResponse({'success': False, 'error': 'សូមជ្រើសរើសរូបភាពយ៉ាងហោចណាស់មួយសន្លឹក'}, status=400)

        from PIL import Image, ImageOps

        orientation = request.POST.get('orientation', 'portrait')
        page_size_choice = request.POST.get('pageSize', 'a4')
        margin_choice = request.POST.get('margin', 'small')
        doc_title = request.POST.get('docTitle', 'SchoolSM_Images_Document').strip() or 'SchoolSM_Images_Document'

        # Page dimensions (A4: 595 x 842 pt, scale x2 for sharp print)
        page_w, page_h = 595, 842
        if page_size_choice == 'letter':
            page_w, page_h = 612, 792

        if orientation == 'landscape':
            page_w, page_h = page_h, page_w

        margin_px = 0
        if margin_choice == 'small':
            margin_px = 25
        elif margin_choice == 'big':
            margin_px = 50

        pdf_pages = []
        for f in images_files:
            img = Image.open(f).convert('RGB')
            img = ImageOps.exif_transpose(img)  # Mobile photo auto-rotation fix

            if page_size_choice == 'fit':
                pdf_pages.append(img)
            else:
                canvas_img = Image.new('RGB', (page_w * 2, page_h * 2), color=(255, 255, 255))
                avail_w = (page_w - (margin_px * 2)) * 2
                avail_h = (page_h - (margin_px * 2)) * 2

                img_ratio = img.width / img.height
                avail_ratio = avail_w / avail_h

                if img_ratio > avail_ratio:
                    new_w = avail_w
                    new_h = int(avail_w / img_ratio)
                else:
                    new_h = avail_h
                    new_w = int(avail_h * img_ratio)

                resized_img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                pos_x = (margin_px * 2) + (avail_w - new_w) // 2
                pos_y = (margin_px * 2) + (avail_h - new_h) // 2
                canvas_img.paste(resized_img, (pos_x, pos_y))
                pdf_pages.append(canvas_img)

        if not pdf_pages:
            return JsonResponse({'success': False, 'error': 'មិនមានទិន្នន័យរូបភាព'}, status=400)

        output_stream = io.BytesIO()
        pdf_pages[0].save(
            output_stream,
            format='PDF',
            save_all=True,
            append_images=pdf_pages[1:],
            resolution=150.0
        )
        output_stream.seek(0)

        response = HttpResponse(output_stream.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{doc_title}.pdf"'
        return response
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'បរាជ័យក្នុងការបង្កើត PDF: {str(e)}'}, status=500)


# ==========================================
# DATABASE BACKUP, SNAPSHOT & RESTORE SUITE
# ==========================================

@login_required
@role_required(['ADMIN'])
def database_backup_view(request):
    """
    Main Database Backup & Snapshot Manager dashboard.
    """
    stats = get_db_statistics()
    backups = list_backups()
    return render(request, 'tools/db_backup.html', {
        'page_title': 'ការគ្រប់គ្រង Database Backup & Snapshot',
        'stats': stats,
        'backups': backups,
        'total_backups': len(backups),
    })


@login_required
@role_required(['ADMIN'])
@require_POST
def api_create_database_backup(request):
    """
    Creates an instant snapshot backup of the current database.
    """
    label = request.POST.get('label', '').strip() or 'Snapshot តាម Web'
    user_info = f"{request.user.get_full_name() or request.user.username} ({request.user.role})"
    try:
        result = create_database_backup(label=label, user_info=user_info)
        messages.success(request, f"បានបង្កើត Backup Snapshot '{result['filename']}' ដោយជោគជ័យ!")
        if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.POST.get('format') == 'json':
            return JsonResponse({'success': True, 'result': result})
        return redirect('tool_database_backup')
    except Exception as e:
        messages.error(request, f"បរាជ័យក្នុងការបង្កើត Backup: {str(e)}")
        if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.POST.get('format') == 'json':
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
        return redirect('tool_database_backup')


@login_required
@role_required(['ADMIN'])
def download_database_backup(request, filename=None):
    """
    Downloads either a specific backup snapshot or the current active db.sqlite3 file.
    """
    if filename == 'current' or not filename:
        db_path = get_db_path()
        if not db_path.exists():
            raise Http404("Database file not found")
        from datetime import datetime
        now_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        response = FileResponse(open(db_path, 'rb'), content_type='application/x-sqlite3')
        response['Content-Disposition'] = f'attachment; filename="school_db_live_{now_str}.sqlite3"'
        return response
    else:
        # Sanitize filename
        safe_filename = os.path.basename(filename)
        backup_file = get_backup_dir() / safe_filename
        if not backup_file.exists():
            raise Http404("Backup snapshot not found")
        response = FileResponse(open(backup_file, 'rb'), content_type='application/x-sqlite3')
        response['Content-Disposition'] = f'attachment; filename="{safe_filename}"'
        return response


@login_required
@role_required(['ADMIN'])
@require_POST
def api_restore_database_backup(request):
    """
    Restores the database from a specified snapshot file.
    """
    filename = request.POST.get('filename', '').strip()
    if not filename:
        messages.error(request, "សូមជ្រើសរើសឯកសារ Backup ដែលចង់ Restore!")
        return redirect('tool_database_backup')

    safe_filename = os.path.basename(filename)
    user_info = f"{request.user.get_full_name() or request.user.username}"
    try:
        result = restore_database_backup(safe_filename, user_info=user_info)
        messages.success(request, f"{result['message']} (ទិន្នន័យមុន Restore ត្រូវបាន Save ទុកក្នុង Safety Backup រួចរាល់)")
        return redirect('tool_database_backup')
    except Exception as e:
        messages.error(request, f"បរាជ័យក្នុងការ Restore Database: {str(e)}")
        return redirect('tool_database_backup')


@login_required
@role_required(['ADMIN'])
@require_POST
def api_upload_restore_database(request):
    """
    Upload an external .sqlite3 database file and restore it as the active database.
    """
    if 'db_file' not in request.FILES:
        messages.error(request, "សូមជ្រើសរើសឯកសារ .sqlite3 ដើម្បី Upload!")
        return redirect('tool_database_backup')

    uploaded_file = request.FILES['db_file']
    if not uploaded_file.name.endswith(('.sqlite3', '.db', '.sqlite')):
        messages.error(request, "ឯកសារត្រូវតែជាប្រភេទ SQLite (.sqlite3 / .db / .sqlite)!")
        return redirect('tool_database_backup')

    # Save uploaded file to backups directory first
    from datetime import datetime
    backup_dir = get_backup_dir()
    now_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    clean_orig_name = "".join(c for c in uploaded_file.name if c.isalnum() or c in ('.', '_', '-'))
    saved_filename = f"db_backup_{now_str}_uploaded_{clean_orig_name}"
    target_path = backup_dir / saved_filename

    with open(target_path, 'wb+') as dest:
        for chunk in uploaded_file.chunks():
            dest.write(chunk)

    # Now restore from this saved file
    user_info = f"{request.user.get_full_name() or request.user.username} (Upload)"
    try:
        result = restore_database_backup(saved_filename, user_info=user_info)
        messages.success(request, f"បាន Upload និង Restore Database ពី {uploaded_file.name} ដោយជោគជ័យ!")
    except Exception as e:
        messages.error(request, f"បរាជ័យក្នុងការ Restore ពី Uploaded File: {str(e)}")

    return redirect('tool_database_backup')


@login_required
@role_required(['ADMIN'])
@require_POST
def api_delete_database_backup(request, filename):
    """
    Deletes a specified backup snapshot file.
    """
    safe_filename = os.path.basename(filename)
    try:
        delete_backup(safe_filename)
        messages.success(request, f"បានលុប Snapshot '{safe_filename}' ដោយជោគជ័យ!")
    except Exception as e:
        messages.error(request, f"បរាជ័យក្នុងការលុប Snapshot: {str(e)}")
    return redirect('tool_database_backup')

